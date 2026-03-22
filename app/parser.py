from docx import Document


def get_heading_level(style_name):
    if "Heading" in style_name:
        try:
            return int(style_name.split()[-1])
        except:
            return None
    return None


def extract_table(table):
    """Convert table into list of rows"""
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    return table_data


def is_table_caption(text):
    """Detect if paragraph is a table caption"""
    lower = text.lower()
    return lower.startswith("table") and ":" in text


def clean_caption(text):
    """Extract caption text from 'Table X: Caption'"""
    if ":" in text:
        return text.split(":", 1)[1].strip()
    return text


def parse_docx(file_path):
    doc = Document(file_path)

    data = {
        "title": "",
        "abstract": [],
        "keywords": [],
        "sections": []
    }

    stack = []
    mode = None

    # Track last caption found (to attach to next table)
    last_caption = None

    elements = list(doc.element.body)

    para_index = 0
    table_index = 0

    for el in elements:

        # ------------------ PARAGRAPH ------------------
        if el.tag.endswith('p'):
            para = doc.paragraphs[para_index]
            para_index += 1

            text = para.text.strip()
            if not text:
                continue

            # Title
            if not data["title"]:
                data["title"] = text
                continue

            lower = text.lower()

            # Abstract
            if lower == "abstract":
                mode = "abstract"
                continue

            # Keywords
            if "keywords" in lower:
                mode = "keywords"
                continue

            # Detect table caption
            if is_table_caption(text):
                last_caption = clean_caption(text)
                continue

            level = get_heading_level(para.style.name)

            # Section heading
            if level:
                mode = "section"

                section = {
                    "heading": text,
                    "level": level,
                    "content": [],
                    "subsections": []
                }

                if level == 1:
                    data["sections"].append(section)
                    stack = [section]
                else:
                    parent = stack[level - 2]
                    parent["subsections"].append(section)

                    if len(stack) >= level:
                        stack[level - 1] = section
                    else:
                        stack.append(section)

            else:
                if mode == "abstract":
                    data["abstract"].append(text)

                elif mode == "keywords":
                    data["keywords"].append(text)

                elif mode == "section" and stack:
                    stack[-1]["content"].append({
                        "type": "text",
                        "value": text
                    })

        # ------------------ TABLE ------------------
        elif el.tag.endswith('tbl'):
            table = doc.tables[table_index]
            table_index += 1

            table_data = extract_table(table)

            if mode == "section" and stack:
                stack[-1]["content"].append({
                    "type": "table",
                    "value": table_data,
                    "caption": last_caption  # may be None
                })

            # reset caption after use
            last_caption = None

    return data