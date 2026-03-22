from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 🔢 Global counter for tables
table_counter = 1


def to_roman(num):
    val = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
    syms = ["M","CM","D","CD","C","XC","L","XL","X","IX","V","IV","I"]
    roman = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman += syms[i]
            num -= val[i]
        i += 1
    return roman


def add_table_caption(doc, caption_text):
    global table_counter
    roman = to_roman(table_counter)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(f"TABLE {roman}")
    run1.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(caption_text.upper())
    run2.bold = True

    table_counter += 1


def add_table(doc, table_data, caption=None):
    if not caption:
        caption = "Sample Table Showing Data"
    add_table_caption(doc, caption)

    rows = len(table_data)
    cols = len(table_data[0]) if rows > 0 else 0
    table = doc.add_table(rows=rows, cols=cols)

    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell

    table.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section(doc, section, level=1):
    doc.add_heading(section["heading"], level=level)

    for item in section["content"]:
        if isinstance(item, dict):
            if item["type"] == "text":
                doc.add_paragraph(item["value"])
            elif item["type"] == "table":
                add_table(doc, item["value"], item.get("caption", None))
        else:
            doc.add_paragraph(item)

    for sub in section["subsections"]:
        add_section(doc, sub, level=level + 1)


def _set_cell_paragraph(cell, text, bold=False, italic=False, font_size=10):
    """Write a line of text into a table cell paragraph, centered."""
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)


def _add_cell_line(cell, text, bold=False, italic=False, font_size=10):
    """Add a new paragraph line inside a cell."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)


def _remove_cell_borders(table):
    """Make all table cell borders invisible."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _insert_element_before(anchor, element):
    """Insert an XML element immediately before anchor."""
    anchor.addprevious(element)


def insert_authors_at_placeholder(doc, placeholder_para, authors):
    """
    Replace {{AUTHORS}} placeholder with IEEE-style author block:
    - 1 author  → single centered block
    - 2–3 authors → side by side in one row (borderless table)
    - 4–6 authors → two rows of up to 3
    - etc.
    Authors sharing the same (org, city, country) are merged into one block.
    Layout per block: Names (bold) → Organization (italic) → City, Country
    """
    anchor = placeholder_para._element
    parent = anchor.getparent()

    if not authors:
        parent.remove(anchor)
        return

    # ── 1. Group authors by (org, location) preserving order ──
    groups = {}
    order = []
    for a in authors:
        org      = a.get('org', '').strip()
        city     = a.get('city', '').strip()
        country  = a.get('country', '').strip()
        location = ", ".join(filter(None, [city, country]))
        key      = (org, location)
        if key not in groups:
            groups[key] = []
            order.append(key)
        full_name = f"{a.get('first', '')} {a.get('last', '')}".strip()
        if full_name:
            groups[key].append(full_name)

    # ── 2. Single author/group → simple centered paragraphs ──
    if len(order) == 1:
        key = order[0]
        org, location = key
        names = groups[key]

        # Build plain XML paragraphs and insert in-place
        new_paras = []
        if names:
            new_paras.append((", ".join(names), True, False, 11))
        if org:
            new_paras.append((org, False, True, 10))
        if location:
            new_paras.append((location, False, False, 10))

        _insert_para_list_at(anchor, parent, new_paras)
        parent.remove(anchor)
        return

    # ── 3. Multiple groups → borderless table, max 3 columns per row ──
    COLS = 3
    rows_needed = (len(order) + COLS - 1) // COLS

    # We need to insert the table at the placeholder position.
    # Strategy: add table normally (appends to end), then move its XML to the right place.

    for row_idx in range(rows_needed):
        row_groups = order[row_idx * COLS : row_idx * COLS + COLS]
        num_cols = len(row_groups)

        tbl = doc.add_table(rows=1, cols=num_cols)
        _remove_cell_borders(tbl)

        # Set equal column widths (~6 inches total / num_cols)
        total_width = Inches(6.0)
        col_width = total_width // num_cols
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = col_width

        for col_idx, key in enumerate(row_groups):
            org, location = key
            names = groups[key]
            cell = tbl.rows[0].cells[col_idx]

            # Clear default empty paragraph
            for p in cell.paragraphs:
                p.clear()

            first = True
            if names:
                if first:
                    _set_cell_paragraph(cell, ", ".join(names), bold=True, font_size=11)
                    first = False
                else:
                    _add_cell_line(cell, ", ".join(names), bold=True, font_size=11)
            if org:
                if first:
                    _set_cell_paragraph(cell, org, italic=True, font_size=10)
                    first = False
                else:
                    _add_cell_line(cell, org, italic=True, font_size=10)
            if location:
                if first:
                    _set_cell_paragraph(cell, location, font_size=10)
                    first = False
                else:
                    _add_cell_line(cell, location, font_size=10)

        # Move the table XML to sit before the anchor placeholder
        tbl_element = tbl._tbl
        tbl_element.getparent().remove(tbl_element)
        anchor.addprevious(tbl_element)

        # Add a small spacer paragraph between rows of tables (not after last)
        if row_idx < rows_needed - 1:
            spacer = _make_spacer_para()
            anchor.addprevious(spacer)

    parent.remove(anchor)


def _make_spacer_para():
    """Empty paragraph as a spacer between author table rows."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), '12')  # 6pt space
    rPr = OxmlElement('w:rPr')
    rPr.append(sz_el)
    pPr.append(rPr)
    p.append(pPr)
    return p


def _make_centered_para_xml(text, bold=False, italic=False, font_size=10):
    """Build a raw centered <w:p> XML element."""
    p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def _insert_para_list_at(anchor, parent, lines):
    """
    Insert a list of (text, bold, italic, font_size) as XML paragraphs
    immediately before anchor, preserving forward order.
    lines: list of (text, bold, italic, font_size)
    """
    prev = None
    for text, bold, italic, font_size in lines:
        el = _make_centered_para_xml(text, bold=bold, italic=italic, font_size=font_size)
        if prev is None:
            anchor.addprevious(el)
        else:
            prev.addnext(el)
        prev = el


def generate_doc(data, output_path):
    global table_counter
    table_counter = 1

    doc = Document("templates/ieee_template.docx")
    authors = data.get("authors", [])

    for para in doc.paragraphs:
        if "{{TITLE}}" in para.text:
            para.text = data["title"]

        elif "{{AUTHORS}}" in para.text:
            insert_authors_at_placeholder(doc, para, authors)

        elif "{{ABSTRACT}}" in para.text:
            para.text = " ".join(data["abstract"])

        elif "{{KEYWORDS}}" in para.text:
            para.text = ", ".join(data["keywords"])

        elif "{{CONTENT}}" in para.text:
            para.text = ""
            for section in data["sections"]:
                add_section(doc, section, level=1)

    doc.save(output_path)
